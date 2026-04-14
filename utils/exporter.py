"""
Mô-đun Xuất - Hỗ trợ Word (DOCX), TXT, Markdown


"""
import os
import re
import logging
import tempfile
from typing import Tuple, Optional
from datetime import datetime
from locales.i18n import t

logger = logging.getLogger(__name__)

MODULE_ROOT = os.path.dirname(os.path.abspath(__file__))
EXPORT_DIR = os.path.join(MODULE_ROOT, "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


def _sanitize_filename(name: str, max_len: int = 120) -> str:
    """Làm sạch các ký tự không hợp lệ trong tên tệp và giới hạn độ dài"""
    if not name or not name.strip():
        name = "novel"
    safe = re.sub(r'[<>:"/\\|?*]', '_', name).strip()
    if len(safe) > max_len:
        safe = safe[:max_len]
    return safe


def _extract_chapters_from_markdown(text: str) -> list:
    """
    Trích xuất thông tin chương từ văn bản tiểu thuyết định dạng Markdown
    
    Returns:
        [{"title": "...", "content": "..."}, ...]
    """
    # Loại bỏ các thẻ HTML phụ trợ (details/summary/b/i) để parse Regex tiêu đề Markdown chính xác
    text = re.sub(r'</?(details|summary|b|i|br|u|strong|em)[^>]*>', '', text)
    
    chapters = []
    current_chapter = None
    content_lines = []

    # Phát hiện tiêu đề chương tổng quát hơn, hỗ trợ '#', '##', '###' và các cấp độ khác, hỗ trợ các biến thể tiếng Trung và không gian
    header_re = re.compile(r'^(?:# {1,6}\s*)?(Chương\s*\d+\s*[\s\S]*|Chương\s*\d+[\s\S]*|Chương\s*\d+\s*[::\s\--–]?.*)$', re.IGNORECASE)

    for line in text.splitlines():
        if not line:
            # Dòng trống có tác dụng ngăn cách đoạn văn nhưng không kết thúc chương
            if current_chapter:
                content_lines.append('')
            continue

        # Phát hiện tiêu đề chương
        if header_re.match(line.strip()):
            # Lưu chương trước
            if current_chapter:
                current_chapter['content'] = '\n'.join(content_lines).strip()
                chapters.append(current_chapter)

            # Trích xuất văn bản tiêu đề
            title_match = re.search(r'(第\s*\d+\s*章[\s\S]*)', line)
            title = title_match.group(1).strip() if title_match else line.strip()
            current_chapter = {'title': title, 'content': ''}
            content_lines = []
            continue

        # Bỏ qua tiêu đề cấp tệp
        if line.strip().startswith('# '):
            continue

        if current_chapter is None:
            # Nếu bạn chưa gặp tiêu đề chương, hãy đặt nội dung của chương đầu tiên
            current_chapter = {'title': t("exporter.first_chapter"), 'content': ''}
            content_lines = [line]
        else:
            content_lines.append(line)

    # lưu chương cuối
    if current_chapter:
        current_chapter['content'] = '\n'.join(content_lines).strip()
        chapters.append(current_chapter)

    return chapters


def export_to_txt(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    Xuất ra định dạng TXT
    
    Args:
        novel_text: Văn bản tiểu thuyết (định dạng Markdown)
        title: Tiêu đề tiểu thuyết
    
    Returns:
        (Đường dẫn tệp, thông tin trạng thái)
    """
    try:
        if not novel_text.strip():
            return None, t("exporter.no_content")
        
        # Trích xuất chương
        chapters = _extract_chapters_from_markdown(novel_text)
        
        if not chapters:
            return None, t("exporter.no_chapters")
        
        # Tạo nội dung TXT
        txt_content = f"{title}\n\n"
        
        for chapter in chapters:
            txt_content += f"{chapter['title']}\n\n"
            txt_content += f"{chapter['content']}\n\n"
            txt_content += "-" * 80 + "\n\n"
        
        # Lưu tập tin (ghi nguyên tử)
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = os.path.join(EXPORT_DIR, filename)
        try:
            with tempfile.NamedTemporaryFile('w', encoding='utf-8', delete=False, dir=EXPORT_DIR) as tmp:
                tmp.write(txt_content)
                tmp_path = tmp.name
            os.replace(tmp_path, filepath)
        except Exception as e:
            logger.error(f"TXT write failed: {e}")
            return None, t("exporter.export_failed", error=str(e))

        logger.info(f"TXT export success: {filename}")
        return filepath, t("exporter.export_success", filename=filename)
    
    except Exception as e:
        logger.error(f"TXT export failed: {e}")
        return None, t("exporter.export_failed", error=str(e))


def export_to_markdown(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    Xuất ra định dạng Markdown
    
    Args:
        novel_text: Văn bản tiểu thuyết (định dạng Markdown)
        title: Tiêu đề tiểu thuyết
    
    Returns:
        (Đường dẫn tệp, thông tin trạng thái)
    """
    try:
        if not novel_text.strip():
            return None, t("exporter.no_content")
        
        # Thêm siêu dữ liệu
        md_content = f"# {title}\n\n"
        md_content += f"*{t('exporter.generated_at', datetime=datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}*\n\n"
        md_content += "---\n\n"
        md_content += novel_text
        
        # Lưu tập tin (ghi nguyên tử)
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        filepath = os.path.join(EXPORT_DIR, filename)
        try:
            with tempfile.NamedTemporaryFile('w', encoding='utf-8', delete=False, dir=EXPORT_DIR) as tmp:
                tmp.write(md_content)
                tmp_path = tmp.name
            os.replace(tmp_path, filepath)
        except Exception as e:
            logger.error(f"Markdown write failed: {e}")
            return None, t("exporter.export_failed", error=str(e))

        logger.info(f"Markdown export success: {filename}")
        return filepath, t("exporter.export_success", filename=filename)
    
    except Exception as e:
        logger.error(f"Markdown export failed: {e}")
        return None, t("exporter.export_failed", error=str(e))


def export_to_docx(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    Xuất ra định dạng Word (DOCX) - Dàn trang chuyên nghiệp
    
    Args:
        novel_text: Văn bản tiểu thuyết (định dạng Markdown)
        title: Tiêu đề tiểu thuyết
    
    Returns:
        (Đường dẫn tệp, thông tin trạng thái)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        return None, t("exporter.missing_docx")
    
    try:
        if not novel_text.strip():
            return None, t("exporter.no_content")
        
        # Trích xuất chương
        chapters = _extract_chapters_from_markdown(novel_text)
        
        if not chapters:
            return None, t("exporter.no_chapters")
        
        doc = Document()
        
        # Kiểu cấu hình
        style = doc.styles['Normal']
        font = style.font
        font.name = t("exporter.body_font")
        font.size = Pt(12)
        
        # phông chữ tiếng trung
        rPr = style.element.get_or_add_rPr()
        rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), t("exporter.body_font"))
        
        # định dạng đoạn văn
        style.paragraph_format.first_line_indent = Pt(24)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        
        # Thêm tên sách
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.name = t("exporter.title_font")
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Cài đặt phông chữ tiếng Trung
        title_rPr = title_run._element.get_or_add_rPr()
        title_rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), t("exporter.title_font"))
        
        # Thêm thông tin tác giả và ngày tháng
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(t("exporter.generated_date", date=datetime.now().strftime('%Y-%m-%d')))
        info_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Dòng trống
        
        # Thêm chương
        for chapter in chapters:
            # Tiêu đề chương
            chapter_title_para = doc.add_paragraph(chapter['title'])
            chapter_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for run in chapter_title_para.runs:
                run.font.name = t('exporter.title_font')
                run.font.size = Pt(16)
                run.font.bold = True
                run_rPr = run._element.get_or_add_rPr()
                run_rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), t("exporter.title_font"))
            
            doc.add_paragraph()  # Dòng trống
            
            # Nội dung chương - thêm theo đoạn
            paragraphs = chapter['content'].split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip(), style='Normal')
            
            doc.add_paragraph()  # Dòng trống giữa các chương
        
        # Lưu tập tin (ghi nguyên tử)
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(EXPORT_DIR, filename)
        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=EXPORT_DIR)
            os.close(tmp_fd)
            doc.save(tmp_path)
            os.replace(tmp_path, filepath)
        except Exception as e:
            logger.error(f"DOCX write failed: {e}")
            return None, t("exporter.export_failed", error=str(e))

        logger.info(f"DOCX export success: {filename}")
        return filepath, t("exporter.export_success", filename=filename)
    
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        return None, t("exporter.export_failed", error=str(e))


def export_to_html(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    Xuất ra định dạng HTML - Có thể đọc trên trình duyệt
    
    Args:
        novel_text: Văn bản tiểu thuyết (định dạng Markdown)
        title: Tiêu đề tiểu thuyết
    
    Returns:
        (Đường dẫn tệp, thông tin trạng thái)
    """
    try:
        import markdown
    except ImportError:
        return None, t("exporter.missing_markdown")
    
    try:
        if not novel_text.strip():
            return None, t("exporter.no_content")
        
        # Chuyển đổi Markdown sang HTML
        html_content = markdown.markdown(novel_text)
        
        # Được gói gọn dưới dạng tài liệu HTML hoàn chỉnh
        full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Arial', 'Times New Roman', serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.8;
            background-color: #f5f5f5;
            color: #333;
        }}
        h1 {{
            text-align: center;
            font-size: 2.5em;
            margin-bottom: 0.5em;
        }}
        h2 {{
            text-align: center;
            font-size: 1.5em;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            border-bottom: 2px solid #ddd;
            padding-bottom: 0.3em;
        }}
        p {{
            text-align: justify;
            text-indent: 2em;
            margin: 1em 0;
        }}
        .info {{
            text-align: center;
            color: #999;
            font-size: 0.9em;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="info">{t('exporter.generated_at', datetime=datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}</p>
    <hr>
    {html_content}
</body>
</html>"""
        
        # Lưu tập tin (ghi nguyên tử)
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        filepath = os.path.join(EXPORT_DIR, filename)
        try:
            with tempfile.NamedTemporaryFile('w', encoding='utf-8', delete=False, dir=EXPORT_DIR) as tmp:
                tmp.write(full_html)
                tmp_path = tmp.name
            os.replace(tmp_path, filepath)
        except Exception as e:
            logger.error(f"HTML write failed: {e}")
            return None, t("exporter.export_failed", error=str(e))

        logger.info(f"HTML export success: {filename}")
        return filepath, t("exporter.export_success", filename=filename)
    
    except Exception as e:
        logger.error(f"HTML export failed: {e}")
        return None, t("exporter.export_failed", error=str(e))


def list_export_files() -> list:
    """Liệt kê tất cả các tập tin xuất"""
    try:
        files = []
        for filename in os.listdir(EXPORT_DIR):
            filepath = os.path.join(EXPORT_DIR, filename)
            if os.path.isfile(filepath):
                file_size = os.path.getsize(filepath)
                file_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                files.append({
                    'name': filename,
                    'path': filepath,
                    'size': file_size,
                    'time': file_time.strftime('%Y-%m-%d %H:%M:%S')
                })
        
        return sorted(files, key=lambda x: x['time'], reverse=True)
    
    except Exception as e:
        logger.error(f"List export files failed: {e}")
        return []
