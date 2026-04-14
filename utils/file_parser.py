"""
Mô-đun phân tích tệp - Hỗ trợ txt/pdf/epub, có theo dõi tiến trình và xử lý lỗi, hỗ trợ mẫu chương tùy chỉnh


"""
import os
import re
import logging
import tempfile
from typing import Tuple, List, IO
from enum import Enum
from dataclasses import dataclass
from locales.i18n import t

logger = logging.getLogger(__name__)

# không thay đổi
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
MIN_PARAGRAPH_LENGTH = 20  # Độ dài đoạn văn tối thiểu


# Mẫu chương mặc định
CHAPTER_PATTERNS = {
    "default": [
        r'第\s*\d+\s*章[：:\s]*.*',
        r'第\s*\d+\s*章',
        r'Chapter\s*\d+',
    ],
    "compact": [
        r'^\d+\.',
        r'^\d+、',
        r'^\d+\s',
    ],
    "brackets": [
        r'《第\d+章》',
        r'「第\d+章」',
    ],
    "english": [
        r'Chapter\s+\d+[:：\s]*.*',
        r'CHAPTER\s+\d+[:：\s]*.*',
        r'Part\s+\d+',
    ],
    "special": [
        r'【.*第\d+章.*】',
        r'≮.*第\d+章.*≯',
        r'◆.*第\d+章.*◆',
    ],
}


@dataclass
class ChapterInfo:
    """Thông tin chương"""
    num: int
    title: str
    content: str
    start_pos: int = 0
    end_pos: int = 0


class FileType(Enum):
    """Loại tệp"""
    TXT = "txt"
    PDF = "pdf"
    EPUB = "epub"
    MD = "md"
    DOCX = "docx"
    UNKNOWN = "unknown"


def get_file_type(file_path: str) -> FileType:
    """Nhận loại tập tin"""
    if not file_path:
        return FileType.UNKNOWN
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".txt":
        return FileType.TXT
    elif ext == ".pdf":
        return FileType.PDF
    elif ext == ".epub":
        return FileType.EPUB
    elif ext == ".md":
        return FileType.MD
    elif ext == ".docx":
        return FileType.DOCX
    else:
        return FileType.UNKNOWN


def parse_txt_file(file_path: str) -> Tuple[List[str], str]:
    """
    Phân tích tệp TXT
    
    Returns:
        (Danh sách đoạn văn, Thông tin trạng thái)
    """
    try:
        # Hỗ trợ truyền vào các đối tượng hoặc đường dẫn tệp
        if hasattr(file_path, 'read'):
            fobj: IO = file_path
            # Cố gắng lấy thuộc tính kích thước
            try:
                fobj.seek(0, os.SEEK_END)
                file_size = fobj.tell()
                fobj.seek(0)
            except Exception:
                file_size = 0
        else:
            file_size = os.path.getsize(file_path)

        if file_size and file_size > MAX_FILE_SIZE:
            return [], t("file_parser.file_too_large", size=f"{file_size / 1024 / 1024:.1f}")

        paragraphs: List[str] = []
        buf_lines: List[str] = []
        total_chars = 0

        # Đọc từng dòng để giảm áp lực bộ nhớ
        if hasattr(file_path, 'read'):
            stream = file_path
        else:
            stream = open(file_path, 'r', encoding='utf-8', errors='ignore')

        try:
            for line in stream:
                stripped = line.rstrip('\n')
                total_chars += len(stripped)

                if stripped.strip() == '':
                    # Dòng trống -> cuối đoạn
                    if buf_lines:
                        para = '\n'.join(buf_lines).strip()
                        if len(para) >= MIN_PARAGRAPH_LENGTH:
                            paragraphs.append(para)
                        buf_lines = []
                    continue

                # hàng thông thường
                buf_lines.append(stripped)

            # đoạn cuối
            if buf_lines:
                para = '\n'.join(buf_lines).strip()
                if len(para) >= MIN_PARAGRAPH_LENGTH:
                    paragraphs.append(para)

        finally:
            if not hasattr(file_path, 'read'):
                stream.close()

        logger.info(f"TXT parse done: {len(paragraphs)} paragraphs")
        return paragraphs, t("file_parser.parse_complete", count=len(paragraphs), chars=total_chars)
    
    except Exception as e:
        logger.error(f"TXT parse failed: {e}")
        return [], t("file_parser.read_failed", error=str(e))


def parse_pdf_file(file_path: str) -> Tuple[List[str], str]:
    """
    Phân tích tệp PDF
    
    Returns:
        (Danh sách đoạn văn, Thông tin trạng thái)
    """
    try:
        import fitz
    except ImportError:
        return [], t("file_parser.missing_pymupdf")
    
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            return [], t("file_parser.file_too_large", size=f"{file_size / 1024 / 1024:.1f}")
        
        text_parts = []
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            try:
                page_text = page.get_text("text")
                text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"PDF page {page_num} parse failed: {e}")
        
        doc.close()
        text = "\n".join(text_parts)
        
        paragraphs = _split_paragraphs(text)
        logger.info(f"PDF parse done: {len(paragraphs)} paragraphs")
        return paragraphs, t("file_parser.parse_complete", count=len(paragraphs), chars=len(text))
    
    except Exception as e:
        logger.error(f"PDF parse failed: {e}")
        return [], t("file_parser.read_failed", error=str(e))


def parse_epub_file(file_path: str) -> Tuple[List[str], str]:
    """
    Phân tích tệp EPUB
    
    Returns:
        (Danh sách đoạn văn, Thông tin trạng thái)
    """
    try:
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError:
        return [], t("file_parser.missing_ebooklib")
    
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            return [], t("file_parser.file_too_large", size=f"{file_size / 1024 / 1024:.1f}")
        
        text_parts = []
        book = epub.read_epub(file_path)
        
        for item in book.get_items():
            if item.get_type() == epub.ITEM_DOCUMENT:
                try:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text_parts.append(soup.get_text(separator="\n"))
                except Exception as e:
                    logger.warning(f"EPUB chapter parse failed: {e}")
        
        text = "\n".join(text_parts)
        paragraphs = _split_paragraphs(text)
        logger.info(f"EPUB parse done: {len(paragraphs)} paragraphs")
        return paragraphs, t("file_parser.parse_complete", count=len(paragraphs), chars=len(text))
    
    except Exception as e:
        logger.error(f"EPUB parse failed: {e}")
        return [], t("file_parser.read_failed", error=str(e))


def parse_md_file(file_path: str) -> Tuple[List[str], str]:
    """
    Phân tích tệp Markdown
    
    Returns:
        (Danh sách đoạn văn, Thông tin trạng thái)
    """
    try:
        # Hỗ trợ truyền vào các đối tượng hoặc đường dẫn tệp
        if hasattr(file_path, 'read'):
            fobj: IO = file_path
            # Cố gắng lấy thuộc tính kích thước
            try:
                fobj.seek(0, os.SEEK_END)
                file_size = fobj.tell()
                fobj.seek(0)
            except Exception:
                file_size = 0
        else:
            file_size = os.path.getsize(file_path)

        if file_size and file_size > MAX_FILE_SIZE:
            return [], t("file_parser.file_too_large", size=f"{file_size / 1024 / 1024:.1f}")

        paragraphs: List[str] = []
        buf_lines: List[str] = []
        total_chars = 0

        # Đọc từng dòng để giảm áp lực bộ nhớ
        if hasattr(file_path, 'read'):
            stream = file_path
        else:
            stream = open(file_path, 'r', encoding='utf-8', errors='ignore')

        try:
            for line in stream:
                stripped = line.rstrip('\n')
                total_chars += len(stripped)

                if stripped.strip() == '':
                    # Dòng trống -> cuối đoạn
                    if buf_lines:
                        para = '\n'.join(buf_lines).strip()
                        if len(para) >= MIN_PARAGRAPH_LENGTH:
                            paragraphs.append(para)
                        buf_lines = []
                    continue

                # hàng thông thường
                buf_lines.append(stripped)

            # đoạn cuối
            if buf_lines:
                para = '\n'.join(buf_lines).strip()
                if len(para) >= MIN_PARAGRAPH_LENGTH:
                    paragraphs.append(para)

        finally:
            if not hasattr(file_path, 'read'):
                stream.close()

        logger.info(f"Markdown parse done: {len(paragraphs)} paragraphs")
        return paragraphs, t("file_parser.parse_complete", count=len(paragraphs), chars=total_chars)
    
    except Exception as e:
        logger.error(f"Markdown parse failed: {e}")
        return [], t("file_parser.read_failed", error=str(e))


def parse_docx_file(file_path: str) -> Tuple[List[str], str]:
    """
    Phân tích tệp tài liệu Word
    
    Returns:
        (Danh sách đoạn văn, Thông tin trạng thái)
    """
    try:
        from docx import Document
    except ImportError:
        return [], t("file_parser.missing_docx")
    
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            return [], t("file_parser.file_too_large", size=f"{file_size / 1024 / 1024:.1f}")
        
        doc = Document(file_path)
        paragraphs: List[str] = []
        total_chars = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) >= MIN_PARAGRAPH_LENGTH:
                paragraphs.append(text)
                total_chars += len(text)
        
        logger.info(f"Word parse done: {len(paragraphs)} paragraphs")
        return paragraphs, t("file_parser.parse_complete", count=len(paragraphs), chars=total_chars)
    
    except Exception as e:
        logger.error(f"Word parse failed: {e}")
        return [], t("file_parser.read_failed", error=str(e))


def parse_novel_file(file_path: str) -> Tuple[List[str], str]:
    """
    Phân tích tệp tiểu thuyết (tự động nhận dạng định dạng)
    
    Args:
        file_path: Đường dẫn tệp
    
    Returns:
        (Danh sách đoạn văn, Thông tin trạng thái)
    """
    if not file_path:
        return [], t("file_parser.no_file")
    
    # Xử lý các đối tượng tệp hoặc luồng tệp được tải lên bởi Gradio
    temp_path = None
    if hasattr(file_path, 'name') and isinstance(file_path.name, str) and os.path.exists(file_path.name):
        file_path = file_path.name
    elif hasattr(file_path, 'read'):
        # Ghi luồng đã tải lên vào một tệp tạm thời để các thư viện xuôi dòng xử lý (PDF/EPUB/DOCX yêu cầu Đường dẫn tệp)
        try:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.tmp')
            chunk = file_path.read(8192)
            while chunk:
                if isinstance(chunk, str):
                    tmp.write(chunk.encode('utf-8'))
                else:
                    tmp.write(chunk)
                chunk = file_path.read(8192)
            tmp.close()
            temp_path = tmp.name
            file_path = temp_path
        except Exception as e:
            logger.error(f"Upload file processing failed: {e}")
            return [], t("file_parser.upload_read_failed", error=str(e))
    
    if not os.path.exists(file_path):
        return [], t("file_parser.file_not_exist", path=file_path)
    
    file_type = get_file_type(file_path)
    
    if file_type == FileType.TXT:
        try:
            return parse_txt_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    elif file_type == FileType.PDF:
        try:
            return parse_pdf_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    elif file_type == FileType.EPUB:
        try:
            return parse_epub_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    elif file_type == FileType.MD:
        try:
            return parse_md_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    elif file_type == FileType.DOCX:
        try:
            return parse_docx_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    else:
        return [], t("file_parser.unsupported_format")


def _split_paragraphs(text: str, min_length: int = MIN_PARAGRAPH_LENGTH) -> List[str]:
    """
    Chia văn bản thành các đoạn văn
    
    Args:
        text: Văn bản gốc
        min_length: Độ dài đoạn văn tối thiểu
    
    Returns:
        Danh sách đoạn văn
    """
    # Chia theo nhiều dòng mới
    raw_paragraphs = re.split(r'\n\s*\n+', text)
    
    # Làm sạch và lọc
    paragraphs = []
    for para in raw_paragraphs:
        para = para.strip()
        # Xóa các điểm đánh dấu đặc biệt như tiêu đề chương
        para = re.sub(r'^(第\d+章|Chapter \d+|第 \d+ 章)[：:：]?\s*', '', para)
        para = re.sub(r'^\s*\*+\s*|\s*\*+\s*$', '', para)
        
        if len(para) >= min_length:
            paragraphs.append(para)
    
    return paragraphs


def estimate_word_count(text: str) -> int:
    """Số lượng ký tự tiếng Trung ước tính (ước tính sơ bộ)"""
    chinese_count = len(re.findall(r'[\u4e00-\u9fff]', text))
    english_count = len(re.findall(r'\b[a-zA-Z]+\b', text))
    # Tiếng Trung được tính là 1 ký tự, tiếng Anh được tính là 0,5 ký tự
    return chinese_count + int(english_count * 0.5)


def parse_novel_by_chapters(
    file_path: str,
    pattern_name: str = "default",
    custom_pattern: str = ""
) -> Tuple[List[ChapterInfo], str]:
    """
    Phân tích tệp tiểu thuyết theo chương

    Args:
        file_path: Đường dẫn tệp
        pattern_name: Tên mẫu định sẵn
        custom_pattern: Biểu thức chính quy tùy chỉnh (nếu được cung cấp, sẽ được ưu tiên sử dụng)

    Returns:
        (Danh sách chương, Thông tin trạng thái)
    """
    try:
        # đọc văn bản
        file_type = get_file_type(file_path)

        if file_type == FileType.TXT:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif file_type == FileType.PDF:
            import fitz
            text_parts = []
            doc = fitz.open(file_path)
            for page in doc:
                text_parts.append(page.get_text("text"))
            doc.close()
            text = "\n".join(text_parts)
        elif file_type == FileType.EPUB:
            from ebooklib import epub
            from bs4 import BeautifulSoup
            text_parts = []
            book = epub.read_epub(file_path)
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text_parts.append(soup.get_text(separator="\n"))
            text = "\n".join(text_parts)
        elif file_type == FileType.MD:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif file_type == FileType.DOCX:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            return [], t("file_parser.unsupported_chapter_format")

        # Xác định biểu thức chính quy để sử dụng
        if custom_pattern and custom_pattern.strip():
            patterns = [custom_pattern.strip()]
        elif pattern_name in CHAPTER_PATTERNS:
            patterns = CHAPTER_PATTERNS[pattern_name]
        else:
            patterns = CHAPTER_PATTERNS.get("default", list(CHAPTER_PATTERNS.values())[0])

        # Tìm tất cả các tiêu đề chương
        chapters = []
        lines = text.split('\n')

        current_chapter_num = 0
        current_chapter_title = ""
        current_chapter_content = []
        chapter_start_pos = 0

        for i, line in enumerate(lines):
            line_stripped = line.strip()
            is_chapter_header = False

            # Kiểm tra xem có mẫu chương nào khớp không
            for pattern in patterns:
                if re.match(pattern, line_stripped, re.IGNORECASE):
                    is_chapter_header = True
                    break

            if is_chapter_header:
                # Lưu chương trước
                if current_chapter_num > 0:
                    content = '\n'.join(current_chapter_content).strip()
                    if content:
                        chapters.append(ChapterInfo(
                            num=current_chapter_num,
                            title=current_chapter_title,
                            content=content,
                            start_pos=chapter_start_pos,
                            end_pos=i
                        ))

                # Trích xuất số chương và tiêu đề
                current_chapter_num += 1
                current_chapter_title = line_stripped
                current_chapter_content = []
                chapter_start_pos = i
            else:
                # Bỏ qua dòng trống nhưng giữ nguyên nội dung
                if line_stripped or current_chapter_content:
                    current_chapter_content.append(line)

        # lưu chương cuối
        if current_chapter_num > 0 and current_chapter_content:
            content = '\n'.join(current_chapter_content).strip()
            if content:
                chapters.append(ChapterInfo(
                    num=current_chapter_num,
                    title=current_chapter_title,
                    content=content,
                    start_pos=chapter_start_pos,
                    end_pos=len(lines)
                ))

        logger.info(f"Chapter parse done: {len(chapters)} chapters")
        return chapters, t("file_parser.chapter_parse_complete", count=len(chapters))

    except Exception as e:
        logger.error(f"Chapter parse failed: {e}")
        return [], t("file_parser.chapter_parse_failed", error=str(e))


def parse_novel_with_custom_template(
    file_path: str,
    custom_template: str
) -> Tuple[List[ChapterInfo], str]:
    """
    Sử dụng mẫu tùy chỉnh để phân tích tiểu thuyết

    Args:
        file_path: Đường dẫn tệp
        custom_template: Mẫu chương tùy chỉnh (hỗ trợ chỗ dành sẵn placeholders)
                       Ví dụ: "Chương {n} {title}" hoặc "Chapter {n}: {title}"

    Returns:
        (Danh sách chương, Thông tin trạng thái)
    """
    if not custom_template or not custom_template.strip():
        return parse_novel_by_chapters(file_path, "default", "")

    # Chuyển đổi mẫu thành biểu thức chính quy
    # {n} hoặc {num} -> (\d+)
    # {title} -> (.*)
    pattern = custom_template.strip()
    pattern = re.escape(pattern)
    pattern = pattern.replace(r'\{n\}', r'(\d+)')
    pattern = pattern.replace(r'\{num\}', r'(\d+)')
    pattern = pattern.replace(r'\{title\}', r'(.*)')
    pattern = pattern.replace(r'\{.*?\}', r'.*')  # Các placeholder khác

    # Đảm bảo khớp với đầu dòng
    if not pattern.startswith('^'):
        pattern = '^' + pattern

    return parse_novel_by_chapters(file_path, custom_pattern=pattern)


def split_by_word_count(text: str, word_count: int) -> List[str]:
    """
    Chia đoạn theo số chữ

    Args:
        text: Văn bản gốc
        word_count: Số chữ mỗi đoạn

    Returns:
        Danh sách văn bản sau khi chia đoạn
    """
    if not text or not text.strip():
        return []

    if word_count <= 0:
        raise ValueError(t("file_parser.word_count_positive"))

    # Chia đều cho số từ
    segments = []
    total_length = len(text)
    start = 0

    while start < total_length:
        end = start + word_count
        if end > total_length:
            end = total_length

        segment = text[start:end].strip()
        if segment:
            segments.append(segment)

        start = end

    logger.info(f"Word count split done: {len(segments)} segments, ~{word_count} each")
    return segments


def split_by_pattern(text: str, pattern: str, keep_marker: bool = True) -> List[str]:
    """
    Chia đoạn theo văn bản/biến cố định

    Args:
        text: Văn bản gốc
        mẫu: Đánh dấu đoạn (Biến hỗ trợ: % Chương (Chương), % Phần (Tiết), % Quay lại (Hồi), hoặc văn bản tùy chỉnh)
        keep_marker: Có giữ lại đánh dấu chia đoạn không

    Returns:
        Danh sách văn bản sau khi chia đoạn
    """
    if not text or not text.strip():
        return []

    if not pattern or not pattern.strip():
        raise ValueError(t("file_parser.split_pattern_empty"))

    # Nhận dạng thông minh: Nếu người dùng nhập "Chương x", "Chương X", v.v., nó sẽ tự động được chuyển đổi thành biểu thức chính quy
    # Kiểm tra xem nó có chứa sự kết hợp của "chương" và "chương", "phần" và "trở lại" không
    pattern_lower = pattern.strip().lower()
    
    # Kiểm tra xem đó có phải là chế độ đơn giản hóa hay không (chẳng hạn như "Chương x", "Chương X")
    if pattern_lower in ['第x章', '第x章', '第x章', '第x章']:
        # Hỗ trợ cả chữ số Trung Quốc và chữ số Ả Rập, sử dụng + để đảm bảo khớp ít nhất một chữ số
        # Sử dụng cái nhìn phủ định để đảm bảo rằng "Chương x" không thể được theo sau bởi các ký tự tiếng Trung (ngoại trừ dấu cách và dấu chấm câu)
        # Định dạng phù hợp: Chương x, Chương x:, Chương x:, Chương x (dấu cách), Chương x (ngắt dòng sau dấu cách)
        # Hỗ trợ định dạng Markdown: ## Chương x
        # Nhưng nó không khớp: đây là chương đầu tiên, nội dung chương đầu tiên, v.v. (có chữ Hán sau đó)
        regex_pattern = r'^[\s# )'
        logger.info("Detected chapter pattern, auto-converting to regex")
    elif pattern_lower in ['第x节', '第x节', '第x节', '第x节']:
        regex_pattern = r'^\s*第\s*[一二三四五六七八九十百千万零〇0123456789]+\s*节\s*[:：\s]*(?![\u4e00-\u9fff])'
        logger.info("Detected section pattern, auto-converting to regex")
    elif pattern_lower in ['第x回', '第x回', '第x回', '第x回']:
        regex_pattern = r'^\s*第\s*[一二三四五六七八九十百千万零〇0123456789]+\s*回\s*[:：\s]*(?![\u4e00-\u9fff])'
        logger.info("Detected episode pattern, auto-converting to regex")
    elif '%章' in pattern_lower or '%节' in pattern_lower or '%回' in pattern_lower:
        # Sử dụng thay thế biến
        regex_pattern = pattern.strip()
        # %Chương -> Khớp "Chương X", "Chương x", v.v. (hỗ trợ chữ số Trung Quốc và Ả Rập)
        regex_pattern = regex_pattern.replace('%章', r'[一二三四五六七八九十百千万零〇0123456789]+\s*章')
        # %Phần -> Khớp "Phần X", "Phần x", v.v. (hỗ trợ chữ số Trung Quốc và Ả Rập)
        regex_pattern = regex_pattern.replace('%节', r'[一二三四五六七八九十百千万零〇0123456789]+\s*节')
        # %chapter -> Khớp "chương X", "chương x", v.v. (hỗ trợ chữ số Trung Quốc và Ả Rập)
        regex_pattern = regex_pattern.replace('%回', r'[一二三四五六七八九十百千万零〇0123456789]+\s*回')
        # Đảm bảo biểu thức chính quy bắt đầu bằng ^ (khớp với đầu dòng)
        if not regex_pattern.startswith('^'):
            regex_pattern = '^' + regex_pattern
    else:
        # Không chứa đánh dấu chương, sử dụng trực tiếp chế độ gốc
        regex_pattern = pattern.strip()

    # Hãy thử chia theo mẫu
    try:
        # Nếu mã thông báo được giữ lại, hãy sử dụng biểu thức chính quy để tìm tất cả các vị trí phù hợp
        if keep_marker:
            # Tìm tất cả các vị trí phù hợp
            matches = list(re.finditer(regex_pattern, text, flags=re.MULTILINE | re.IGNORECASE))

            if not matches:
                # Không khớp, trả lại toàn bộ văn bản
                logger.warning(f"No pattern match: {regex_pattern}, returning full text")
                return [text.strip()] if text.strip() else []

            segments = []
            prev_end = 0

            for match in matches:
                # Nhận thẻ phù hợp
                marker = match.group(0)

                # Lấy nội dung trước dấu (nếu có)
                if prev_end < match.start():
                    prev_content = text[prev_end:match.start()].strip()
                    if prev_content:
                        segments.append(prev_content)

                # Thêm thẻ
                segments.append(marker.strip())

                prev_end = match.end()

            # Thêm đoạn cuối
            if prev_end < len(text):
                last_content = text[prev_end:].strip()
                if last_content:
                    segments.append(last_content)

            # Hợp nhất đánh dấu và nội dung
            result = []
            i = 0
            while i < len(segments):
                # Nếu nó hiện là một nhãn hiệu và có nội dung đằng sau nó
                if i + 1 < len(segments):
                    result.append((segments[i] + segments[i + 1]).strip())
                    i += 2
                else:
                    # chỉ đánh dấu hoặc nội dung
                    if segments[i].strip():
                        result.append(segments[i].strip())
                    i += 1

            segments = result
        else:
            # Không giữ lại điểm đánh dấu và chia trực tiếp
            segments = re.split(regex_pattern, text, flags=re.MULTILINE | re.IGNORECASE)

        # Dọn dẹp các đoạn văn trống
        segments = [seg.strip() for seg in segments if seg.strip()]

        logger.info(f"Pattern split done: {len(segments)} segments")
        return segments

    except re.error as e:
        raise ValueError(t("file_parser.invalid_regex", error=str(e)))
